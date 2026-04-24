"""FastAPI 服务 — 公众号/头条文章生成器 v4（+审计+单节重写+导出Word+文章模板）"""
from __future__ import annotations
import asyncio, json, os, time
from pathlib import Path
from typing import Literal

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel

from llm import LLM
from pipeline import Pipeline, StyleProfile
from validator import Validator

# ── 初始化 ────────────────────────────────────────────────────────────────────

app = FastAPI(title="InkFlow - AI 写作工坊", version="2.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

BASE_DIR = Path(__file__).parent
ENV_PATH = BASE_DIR / ".env"

# ── 并发保护 ──────────────────────────────────────────────────────────────────
# 同一时间只允许一个生成任务运行，防止 API 并发超限
_generation_lock = asyncio.Lock()
_generation_active = {"active": False, "started_at": 0, "task": ""}


def _load_env():
    if ENV_PATH.exists():
        for line in ENV_PATH.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ[k.strip()] = v.strip()


def _create_llm(temperature: float = 0.7) -> LLM:
    _load_env()
    api_key = os.environ.get("DEEPSEEK_API_KEY", "")
    if not api_key:
        raise HTTPException(400, "请先配置 DEEPSEEK_API_KEY")
    return LLM(
        api_key=api_key,
        base_url=os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com/v1"),
        model=os.environ.get("DEEPSEEK_MODEL", "deepseek-chat"),
        temperature=temperature,
    )


def _pipeline(temperature: float = 0.7) -> Pipeline:
    return Pipeline(_create_llm(temperature))


# ── 请求模型 ──────────────────────────────────────────────────────────────────

class GenerateReq(BaseModel):
    topic: str
    platform: Literal["wechat", "toutiao"] = "wechat"
    mode: Literal["干货型", "争议型", "故事型", "测评型"] = "干货型"
    source_text: str = ""
    style_profile_id: str = ""
    extra_requirements: str = ""

class GenerateOutlineReq(BaseModel):
    topic: str
    platform: Literal["wechat", "toutiao"] = "wechat"
    mode: Literal["干货型", "争议型", "故事型", "测评型"] = "干货型"
    source_text: str = ""
    style_profile_id: str = ""
    extra_requirements: str = ""

class GenerateContentReq(BaseModel):
    outline: dict
    topic: str = ""
    platform: str = "wechat"
    source_text: str = ""
    style_profile_id: str = ""

class GenerateTitlesReq(BaseModel):
    topic: str
    outline_summary: str = ""
    platform: str = "wechat"
    count: int = 10

class UpdateOutlineReq(BaseModel):
    outline: dict

class ExtractStyleReq(BaseModel):
    name: str = "我的风格"

class SaveSettingsReq(BaseModel):
    deepseek_api_key: str = ""
    deepseek_base_url: str = "https://api.deepseek.com/v1"
    deepseek_model: str = "deepseek-chat"


# ── 健康检查 ──────────────────────────────────────────────────────────────────

@app.get("/api/health")
def health_check():
    _load_env()
    configured = bool(os.environ.get("DEEPSEEK_API_KEY", ""))
    return {
        "status": "ok",
        "version": "2.0.0",
        "configured": configured,
        "generation_active": _generation_active["active"],
        "uptime": time.time(),
    }


# ── 页面路由 ──────────────────────────────────────────────────────────────────

@app.get("/")
def index():
    return FileResponse(str(BASE_DIR / "index.html"))


# ── 设置 ──────────────────────────────────────────────────────────────────────

@app.get("/api/settings")
def get_settings():
    _load_env()
    key = os.environ.get("DEEPSEEK_API_KEY", "")
    return {
        "deepseek_api_key": key[:8] + "***" if key and len(key) > 8 else key,
        "deepseek_base_url": os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com/v1"),
        "deepseek_model": os.environ.get("DEEPSEEK_MODEL", "deepseek-chat"),
        "configured": bool(key),
    }


@app.post("/api/settings")
def save_settings(req: SaveSettingsReq):
    lines = [
        "# InkFlow 配置",
        f"DEEPSEEK_BASE_URL={req.deepseek_base_url}",
        f"DEEPSEEK_MODEL={req.deepseek_model}",
    ]
    if req.deepseek_api_key and not req.deepseek_api_key.endswith("***"):
        lines.append(f"DEEPSEEK_API_KEY={req.deepseek_api_key}")
    else:
        if ENV_PATH.exists():
            for line in ENV_PATH.read_text(encoding="utf-8").splitlines():
                if line.strip().startswith("DEEPSEEK_API_KEY="):
                    lines.append(line.strip())
    ENV_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return {"ok": True}


# ── 文章列表 ──────────────────────────────────────────────────────────────────

@app.get("/api/articles")
def list_articles():
    return Pipeline.list_articles()


# ── 写作统计 ──────────────────────────────────────────────────────────────────

@app.get("/api/stats")
def get_stats():
    art_dir = BASE_DIR / "data" / "articles"
    if not art_dir.exists():
        return {"total": 0, "total_words": 0, "avg_score": 0, "by_platform": {},
                "by_mode": {}, "by_month": {}, "top_topics": []}

    articles = []
    for f in art_dir.glob("*.json"):
        try:
            articles.append(json.loads(f.read_text(encoding="utf-8")))
        except Exception:
            pass

    if not articles:
        return {"total": 0, "total_words": 0, "avg_score": 0, "by_platform": {},
                "by_mode": {}, "by_month": {}, "top_topics": []}

    total = len(articles)
    total_words = sum(a.get("word_count", 0) for a in articles)
    scores = [a.get("validation", {}).get("score", 0) for a in articles if a.get("validation", {}).get("score")]
    avg_score = round(sum(scores) / len(scores), 1) if scores else 0

    # 按平台
    by_platform = {}
    for a in articles:
        p = a.get("platform", "unknown")
        by_platform[p] = by_platform.get(p, 0) + 1

    # 按类型
    by_mode = {}
    for a in articles:
        m = a.get("mode", "unknown")
        by_mode[m] = by_mode.get(m, 0) + 1

    # 按月份
    by_month = {}
    for a in articles:
        dt = a.get("created_at", "")
        if dt and len(dt) >= 7:
            month = dt[:7]
            by_month[month] = by_month.get(month, 0) + 1

    # 最近文章
    recent = sorted(articles, key=lambda x: x.get("created_at", ""), reverse=True)[:5]
    recent_list = [{"id": a.get("id"), "topic": a.get("topic", ""),
                    "word_count": a.get("word_count", 0),
                    "score": a.get("validation", {}).get("score"),
                    "created_at": a.get("created_at", "")} for a in recent]

    return {
        "total": total,
        "total_words": total_words,
        "avg_score": avg_score,
        "by_platform": by_platform,
        "by_mode": by_mode,
        "by_month": by_month,
        "recent": recent_list,
    }


@app.get("/api/articles/{article_id}")
def get_article(article_id: str):
    try:
        return Pipeline.load_article(article_id)
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))


@app.delete("/api/articles/{article_id}")
def delete_article(article_id: str):
    path = BASE_DIR / "data" / "articles" / f"{article_id}.json"
    if path.exists():
        path.unlink()
    return {"ok": True}


# ── 风格管理 ──────────────────────────────────────────────────────────────────

@app.get("/api/styles")
def list_styles():
    return Pipeline.list_styles()


@app.get("/api/styles/{style_id}")
def get_style(style_id: str):
    try:
        profile = Pipeline.load_style(style_id)
        return {
            "id": profile.id, "name": profile.name,
            "tone": profile.tone, "structure_style": profile.structure_style,
            "vocabulary": profile.vocabulary, "sentence_patterns": profile.sentence_patterns,
            "punctuation_habits": profile.punctuation_habits,
            "sample_summary": profile.sample_summary,
        }
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))


@app.delete("/api/styles/{style_id}")
def delete_style(style_id: str):
    path = BASE_DIR / "data" / "styles" / f"{style_id}.json"
    if path.exists():
        path.unlink()
    return {"ok": True}


@app.post("/api/styles/extract")
async def extract_style(
    name: str = Form("我的风格"),
    files: list[UploadFile] = File(...),
):
    samples = []
    for f in files[:5]:
        content = await f.read()
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("gbk", errors="replace")
        samples.append(text)

    if not samples:
        raise HTTPException(400, "请上传至少 1 个样本文件")

    try:
        profile = await asyncio.to_thread(
            _pipeline(0.3).extract_style, samples, name,
        )
        return {
            "ok": True,
            "profile": {
                "id": profile.id, "name": profile.name,
                "tone": profile.tone, "structure_style": profile.structure_style,
                "vocabulary": profile.vocabulary, "sentence_patterns": profile.sentence_patterns,
                "punctuation_habits": profile.punctuation_habits,
                "sample_summary": profile.sample_summary,
            },
        }
    except Exception as e:
        raise HTTPException(500, f"风格提取失败：{e}")


# ── 素材提取 ──────────────────────────────────────────────────────────────────

class ExtractMaterialReq(BaseModel):
    source_text: str
    topic: str


@app.post("/api/extract-material")
async def extract_material(req: ExtractMaterialReq):
    try:
        material = await asyncio.to_thread(
            _pipeline(0.3).extract_material, req.source_text, req.topic,
        )
        return {"ok": True, "material": material}
    except Exception as e:
        raise HTTPException(500, f"素材提取失败：{e}")


# ── 大纲生成 ──────────────────────────────────────────────────────────────────

@app.post("/api/generate/outline")
async def generate_outline(req: GenerateOutlineReq):
    style = None
    if req.style_profile_id:
        try:
            style = Pipeline.load_style(req.style_profile_id)
        except FileNotFoundError:
            pass

    material = None
    if req.source_text.strip():
        try:
            material = await asyncio.to_thread(
                _pipeline(0.3).extract_material, req.source_text, req.topic,
            )
        except Exception:
            pass

    try:
        outline = await asyncio.to_thread(
            _pipeline(0.8).generate_outline,
            req.topic, req.platform, req.mode, material, style, req.extra_requirements,
        )
        return {
            "ok": True,
            "outline": {
                "title_candidates": outline.title_candidates,
                "selected_title": outline.selected_title,
                "hook": outline.hook,
                "sections": outline.sections,
                "cta": outline.cta,
                "tags": outline.tags,
            },
        }
    except Exception as e:
        raise HTTPException(500, f"大纲生成失败：{e}")


# ── 标题优化 ──────────────────────────────────────────────────────────────────

@app.post("/api/generate/titles")
async def generate_titles(req: GenerateTitlesReq):
    try:
        titles = await asyncio.to_thread(
            _pipeline(0.9).generate_titles,
            req.topic, req.outline_summary, req.platform, req.count,
        )
        return {"ok": True, "titles": titles}
    except Exception as e:
        raise HTTPException(500, f"标题生成失败：{e}")


# ── 正文生成（普通模式）──────────────────────────────────────────────────────

@app.post("/api/generate/content")
async def generate_content(req: GenerateContentReq):
    # 并发检查
    if _generation_active["active"]:
        raise HTTPException(429, f"另一个生成任务正在进行中（{_generation_active['task']}），请等待完成")

    from pipeline import Outline
    outline = Outline(
        title_candidates=req.outline.get("title_candidates", []),
        selected_title=req.outline.get("selected_title", ""),
        hook=req.outline.get("hook", ""),
        sections=req.outline.get("sections", []),
        cta=req.outline.get("cta", ""),
        tags=req.outline.get("tags", []),
    )

    style = None
    if req.style_profile_id:
        try:
            style = Pipeline.load_style(req.style_profile_id)
        except FileNotFoundError:
            pass

    material = None
    if req.source_text.strip():
        try:
            material = await asyncio.to_thread(
                _pipeline(0.3).extract_material, req.source_text, req.topic,
            )
        except Exception:
            pass

    _generation_active.update({"active": True, "started_at": time.time(), "task": "正文生成"})
    try:
        content = await asyncio.to_thread(
            _pipeline(0.8).generate_content,
            req.topic, outline, req.platform, material, style,
        )
        return {"ok": True, "content": content, "word_count": len(content)}
    except Exception as e:
        raise HTTPException(500, f"正文生成失败：{e}")
    finally:
        _generation_active.update({"active": False, "task": ""})


# ── 正文生成（SSE 流式模式）──────────────────────────────────────────────────

@app.post("/api/generate/content/stream")
async def generate_content_stream(req: GenerateContentReq):
    if _generation_active["active"]:
        raise HTTPException(429, f"另一个生成任务正在进行中，请等待完成")

    from pipeline import Outline
    outline = Outline(
        title_candidates=req.outline.get("title_candidates", []),
        selected_title=req.outline.get("selected_title", ""),
        hook=req.outline.get("hook", ""),
        sections=req.outline.get("sections", []),
        cta=req.outline.get("cta", ""),
        tags=req.outline.get("tags", []),
    )

    style = None
    if req.style_profile_id:
        try:
            style = Pipeline.load_style(req.style_profile_id)
        except FileNotFoundError:
            pass

    material = None
    if req.source_text.strip():
        try:
            material = await asyncio.to_thread(
                _pipeline(0.3).extract_material, req.source_text, req.topic,
            )
        except Exception:
            pass

    _generation_active.update({"active": True, "started_at": time.time(), "task": "流式正文生成"})

    def event_stream():
        try:
            pipeline = _pipeline(0.8)
            for event in pipeline.generate_content_stream(
                req.topic, outline, req.platform, material, style
            ):
                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"
        finally:
            _generation_active.update({"active": False, "task": ""})

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ── 修订建议 ──────────────────────────────────────────────────────────────────

class SuggestReq(BaseModel):
    content: str
    platform: str = "wechat"
    outline: dict = {}


@app.post("/api/revise/suggest")
async def suggest_revisions(req: SuggestReq):
    v = Validator()
    result = v.validate(req.content, req.platform)

    suggestions = []
    for issue in result.issues:
        if issue.rule == "AI_MARKER":
            suggestions.append({"text": "减少 AI 痕迹词", "detail": issue.description, "icon": "🤖"})
        elif issue.rule == "FORBIDDEN":
            suggestions.append({"text": "删除禁止句式", "detail": issue.description, "icon": "🚫"})
        elif issue.rule == "LONG_PARA":
            suggestions.append({"text": "拆分过长段落", "detail": issue.description, "icon": "📱"})
        elif issue.rule == "CONSECUTIVE_LE":
            suggestions.append({"text": "减少「了」字连用", "detail": issue.description, "icon": "✍️"})
        elif issue.rule == "REPORT":
            suggestions.append({"text": "去掉报告腔", "detail": issue.description, "icon": "📋"})
        elif issue.rule == "COLLECTIVE":
            suggestions.append({"text": "删除集体套话", "detail": issue.description, "icon": "👥"})
        elif issue.rule == "META":
            suggestions.append({"text": "去掉元叙事词汇", "detail": issue.description, "icon": "📖"})
        elif issue.rule == "AI_FILLER":
            suggestions.append({"text": "精简套话连接词", "detail": issue.description, "icon": "✂️"})
        elif issue.rule == "REPETITIVE":
            suggestions.append({"text": "避免重复表达", "detail": issue.description, "icon": "🔁"})
        elif issue.rule == "EXCLAMATION":
            suggestions.append({"text": "减少感叹号", "detail": issue.description, "icon": "❗"})

    suggestions.append({"text": "开头更抓人", "detail": "让前3句更有冲击力，制造悬念或共鸣", "icon": "🎣"})
    suggestions.append({"text": "结尾加互动引导", "detail": "引导点赞、转发、评论", "icon": "💬"})
    suggestions.append({"text": "语气更口语化", "detail": "像跟朋友聊天一样，去掉书面腔", "icon": "🗣️"})
    suggestions.append({"text": "增加具体案例", "detail": "用真实故事或数据支撑观点", "icon": "📊"})

    section_suggestions = []
    if req.outline and req.outline.get("sections"):
        for sec in req.outline["sections"]:
            title = sec.get("title", "")
            if title:
                section_suggestions.append(title)

    return {
        "ok": True,
        "suggestions": suggestions[:8],
        "sections": section_suggestions,
        "score": result.score,
        "passed": result.passed,
        "issue_count": len(result.issues),
    }


# ── 多轮修订 ──────────────────────────────────────────────────────────────────

class ReviseReq(BaseModel):
    content: str
    instruction: str
    topic: str = ""
    outline: dict = {}
    platform: str = "wechat"
    style_profile_id: str = ""
    section_title: str = ""


@app.post("/api/revise")
async def revise_content(req: ReviseReq):
    if not req.content.strip():
        raise HTTPException(400, "没有可修订的内容")
    if not req.instruction.strip():
        raise HTTPException(400, "请输入修改意见")

    if _generation_active["active"]:
        raise HTTPException(429, f"另一个生成任务正在进行中，请等待完成")

    style = None
    if req.style_profile_id:
        try:
            style = Pipeline.load_style(req.style_profile_id)
        except FileNotFoundError:
            pass

    outline = None
    if req.outline:
        from pipeline import Outline
        outline = Outline(
            title_candidates=req.outline.get("title_candidates", []),
            selected_title=req.outline.get("selected_title", ""),
            hook=req.outline.get("hook", ""),
            sections=req.outline.get("sections", []),
            cta=req.outline.get("cta", ""),
            tags=req.outline.get("tags", []),
        )

    _generation_active.update({"active": True, "started_at": time.time(), "task": "修订"})
    try:
        revised = await asyncio.to_thread(
            _pipeline(0.5).revise_content,
            req.content, req.instruction, req.topic,
            outline, req.platform, style, req.section_title,
        )
        return {"ok": True, "content": revised, "word_count": len(revised)}
    except Exception as e:
        raise HTTPException(500, f"修订失败：{e}")
    finally:
        _generation_active.update({"active": False, "task": ""})


# ── 排版 ──────────────────────────────────────────────────────────────────────

class FormatReq(BaseModel):
    content: str
    title: str = ""
    tags: list[str] = []
    template: str = "minimal"


@app.post("/api/format")
async def format_html(req: FormatReq):
    try:
        html = await asyncio.to_thread(
            _pipeline(0.3).format_html, req.content, req.title, req.tags, req.template,
        )
        return {"ok": True, "html": html}
    except Exception as e:
        import traceback; traceback.print_exc()
        raise HTTPException(500, f"排版失败：{e}")


# ── 排版模板列表 ──────────────────────────────────────────────────────────────

@app.get("/api/format/templates")
def list_format_templates():
    from pipeline import FORMAT_TEMPLATES
    return [{"id": k, "name": v["name"], "desc": v["desc"]} for k, v in FORMAT_TEMPLATES.items()]


# ── 验证 ──────────────────────────────────────────────────────────────────────

class ValidateReq(BaseModel):
    content: str
    platform: str = "wechat"


@app.post("/api/validate")
def validate(req: ValidateReq):
    v = Validator()
    result = v.validate(req.content, req.platform)
    return {
        "passed": result.passed,
        "score": result.score,
        "issues": [{"rule": i.rule, "severity": i.severity,
                     "description": i.description, "excerpt": i.excerpt}
                    for i in result.issues],
    }


# ── 参考链接抓取 ──────────────────────────────────────────────────────────────

class FetchReferencesReq(BaseModel):
    urls: list[str]
    topic: str


@app.post("/api/fetch-references")
async def fetch_references(req: FetchReferencesReq):
    if not req.urls:
        raise HTTPException(400, "请提供至少一个参考链接")
    if len(req.urls) > 5:
        raise HTTPException(400, "最多支持 5 个参考链接")

    try:
        material = await asyncio.to_thread(
            _pipeline(0.3).extract_material_from_urls, req.urls, req.topic,
        )
        errors = material.pop("_fetch_errors", [])
        return {"ok": True, "material": material, "fetch_errors": errors}
    except Exception as e:
        import traceback; traceback.print_exc()
        raise HTTPException(500, f"参考链接抓取失败：{e}")


# ── 图片建议 ──────────────────────────────────────────────────────────────────

class SuggestImagesReq(BaseModel):
    content: str
    outline: dict


@app.post("/api/suggest-images")
async def suggest_images(req: SuggestImagesReq):
    from pipeline import Outline
    outline = Outline(
        title_candidates=req.outline.get("title_candidates", []),
        selected_title=req.outline.get("selected_title", ""),
        hook=req.outline.get("hook", ""),
        sections=req.outline.get("sections", []),
        cta=req.outline.get("cta", ""),
        tags=req.outline.get("tags", []),
    )
    try:
        suggestions = await asyncio.to_thread(
            _pipeline(0.3).suggest_images, req.content, outline,
        )
        return {"ok": True, "suggestions": suggestions}
    except Exception as e:
        raise HTTPException(500, f"图片建议生成失败：{e}")


# ── 批量生成 ──────────────────────────────────────────────────────────────────

class BatchGenerateReq(BaseModel):
    topics: list[str]
    platform: str = "wechat"
    mode: str = "干货型"
    style_profile_id: str = ""


@app.post("/api/generate/batch")
async def generate_batch(req: BatchGenerateReq):
    if not req.topics:
        raise HTTPException(400, "请提供至少一个主题")
    if len(req.topics) > 10:
        raise HTTPException(400, "批量生成最多支持 10 个主题")

    style = None
    if req.style_profile_id:
        try:
            style = Pipeline.load_style(req.style_profile_id)
        except FileNotFoundError:
            pass

    _generation_active.update({"active": True, "started_at": time.time(), "task": f"批量生成{len(req.topics)}篇"})
    try:
        results = await asyncio.to_thread(
            _pipeline(0.8).generate_batch, req.topics, req.platform, req.mode, style,
        )
        return {"ok": True, "results": results}
    except Exception as e:
        raise HTTPException(500, f"批量生成失败：{e}")
    finally:
        _generation_active.update({"active": False, "task": ""})


# ── SEO 优化建议 ──────────────────────────────────────────────────────────────

class SeoReq(BaseModel):
    content: str
    topic: str


@app.post("/api/seo")
async def seo_optimize(req: SeoReq):
    try:
        result = await asyncio.to_thread(
            _pipeline(0.3).optimize_seo, req.content, req.topic,
        )
        return {"ok": True, "seo": result}
    except Exception as e:
        raise HTTPException(500, f"SEO 分析失败：{e}")


# ── 一键生成 ──────────────────────────────────────────────────────────────────

@app.post("/api/generate/full")
async def generate_full(req: GenerateReq):
    if _generation_active["active"]:
        raise HTTPException(429, f"另一个生成任务正在进行中，请等待完成")

    style = None
    if req.style_profile_id:
        try:
            style = Pipeline.load_style(req.style_profile_id)
        except FileNotFoundError:
            pass

    _generation_active.update({"active": True, "started_at": time.time(), "task": "一键完整生成"})

    def _run():
        return _pipeline(0.8).run_full(
            topic=req.topic,
            platform=req.platform,
            mode=req.mode,
            source_text=req.source_text,
            style_profile=style,
            extra_requirements=req.extra_requirements,
        )

    try:
        article = await asyncio.to_thread(_run)
        from pipeline import _to_dict
        return {"ok": True, "article": _to_dict(article)}
    except Exception as e:
        import traceback; traceback.print_exc()
        raise HTTPException(500, f"生成失败：{e}")
    finally:
        _generation_active.update({"active": False, "task": ""})



# ── 单节重新生成 ──────────────────────────────────────────────────────────────

class RegenerateSectionReq(BaseModel):
    outline: dict
    section_index: int
    content: str
    topic: str = ""
    platform: str = "wechat"
    source_text: str = ""
    style_profile_id: str = ""


@app.post("/api/generate/section")
async def regenerate_section(req: RegenerateSectionReq):
    if _generation_active["active"]:
        raise HTTPException(429, "另一个生成任务正在进行中，请等待完成")

    from pipeline import Outline
    outline = Outline(
        title_candidates=req.outline.get("title_candidates", []),
        selected_title=req.outline.get("selected_title", ""),
        hook=req.outline.get("hook", ""),
        sections=req.outline.get("sections", []),
        cta=req.outline.get("cta", ""),
        tags=req.outline.get("tags", []),
    )

    style = None
    if req.style_profile_id:
        try:
            style = Pipeline.load_style(req.style_profile_id)
        except FileNotFoundError:
            pass

    material = None
    if req.source_text.strip():
        try:
            material = await asyncio.to_thread(
                _pipeline(0.3).extract_material, req.source_text, req.topic,
            )
        except Exception:
            pass

    _generation_active.update({"active": True, "started_at": time.time(), "task": "单节重写"})
    try:
        content = await asyncio.to_thread(
            _pipeline(0.8).regenerate_section,
            req.topic, outline, req.section_index, req.content,
            req.platform, style, material,
        )
        return {"ok": True, "content": content, "word_count": len(content)}
    except Exception as e:
        raise HTTPException(500, f"单节重写失败：{e}")
    finally:
        _generation_active.update({"active": False, "task": ""})


# ── 文章综合审计 ──────────────────────────────────────────────────────────────

class AuditReq(BaseModel):
    content: str
    topic: str = ""
    platform: str = "wechat"


@app.post("/api/audit")
async def audit_article(req: AuditReq):
    try:
        result = await asyncio.to_thread(
            _pipeline(0.3).audit_article, req.content, req.topic, req.platform,
        )
        return {"ok": True, "audit": result}
    except Exception as e:
        import traceback; traceback.print_exc()
        raise HTTPException(500, f"审计失败：{e}")


# ── 导出 Word ─────────────────────────────────────────────────────────────────

class ExportDocxReq(BaseModel):
    content: str
    title: str = ""


@app.post("/api/export/docx")
async def export_docx(req: ExportDocxReq):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io, re

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(11)

    if req.title:
        heading = doc.add_heading(req.title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = req.content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('## ') or (line.startswith('**') and line.endswith('**')):
            title_text = line.lstrip('#').strip().strip('*').strip()
            doc.add_heading(title_text, level=2)
        elif line.startswith('# '):
            doc.add_heading(line.lstrip('#').strip(), level=1)
        else:
            para = doc.add_paragraph()
            parts = re.split(r'\*\*(.+?)\*\*', line)
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = para.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    from fastapi.responses import StreamingResponse
    import urllib.parse
    filename = f"{req.title or 'article'}.docx"
    encoded = urllib.parse.quote(filename)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"}
    )


# ── 文章模板列表 ──────────────────────────────────────────────────────────────

@app.get("/api/templates")
def list_article_templates():
    from pipeline import ARTICLE_TEMPLATES
    return [{"id": k, "name": v["name"], "desc": v["desc"], "mode": v["mode"]}
            for k, v in ARTICLE_TEMPLATES.items()]


# ── 保存/更新文章 ─────────────────────────────────────────────────────────────

class SaveArticleReq(BaseModel):
    id: str = ""
    topic: str = ""
    platform: str = "wechat"
    mode: str = "干货型"
    outline: dict = {}
    content: str = ""
    html: str = ""
    source_text: str = ""
    material: dict = {}
    style_profile_id: str = ""
    extra_requirements: str = ""
    revision_history: list = []
    tags: list[str] = []


@app.post("/api/articles/save")
def save_article(req: SaveArticleReq):
    import uuid
    from datetime import datetime, timezone

    article_id = req.id or f"art_{uuid.uuid4().hex[:8]}"
    now = datetime.now(timezone.utc)
    data = {
        "id": article_id,
        "topic": req.topic,
        "platform": req.platform,
        "mode": req.mode,
        "status": "done" if req.content else "outlined",
        "outline": req.outline,
        "content": req.content,
        "html": req.html,
        "source_text": req.source_text,
        "material": req.material,
        "style_profile_id": req.style_profile_id,
        "extra_requirements": req.extra_requirements,
        "word_count": len(req.content),
        "created_at": now.isoformat(),
        "revision_history": req.revision_history,
        "tags": req.tags,
    }

    art_dir = BASE_DIR / "data" / "articles"
    art_dir.mkdir(parents=True, exist_ok=True)
    path = art_dir / f"{article_id}.json"

    # 自动版本管理：保存前先存旧版本
    if path.exists() and req.content:
        try:
            old = json.loads(path.read_text(encoding="utf-8"))
            if old.get("content") and old["content"] != req.content:
                ver_dir = art_dir / article_id / "versions"
                ver_dir.mkdir(parents=True, exist_ok=True)
                ver_path = ver_dir / f"{now.strftime('%Y%m%d_%H%M%S')}.json"
                ver_data = {
                    "content": old["content"],
                    "word_count": old.get("word_count", 0),
                    "timestamp": now.isoformat(),
                    "snapshot": old.get("content", "")[:200],
                }
                ver_path.write_text(json.dumps(ver_data, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass

    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"ok": True, "id": article_id}


# ── 标签管理 ──────────────────────────────────────────────────────────────────

@app.get("/api/tags")
def list_tags():
    art_dir = BASE_DIR / "data" / "articles"
    if not art_dir.exists():
        return []
    tags_count = {}
    for f in art_dir.glob("*.json"):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            for t in data.get("tags", []):
                tags_count[t] = tags_count.get(t, 0) + 1
        except Exception:
            pass
    return [{"name": k, "count": v} for k, v in sorted(tags_count.items(), key=lambda x: -x[1])]


@app.get("/api/articles/tag/{tag}")
def list_articles_by_tag(tag: str):
    art_dir = BASE_DIR / "data" / "articles"
    if not art_dir.exists():
        return []
    articles = []
    for f in sorted(art_dir.glob("*.json"), reverse=True):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            if tag in data.get("tags", []):
                articles.append({
                    "id": data.get("id"),
                    "topic": data.get("topic"),
                    "platform": data.get("platform"),
                    "mode": data.get("mode"),
                    "status": data.get("status"),
                    "word_count": data.get("word_count"),
                    "score": data.get("validation", {}).get("score"),
                    "title": data.get("outline", {}).get("selected_title"),
                    "created_at": data.get("created_at"),
                    "tags": data.get("tags", []),
                })
        except Exception:
            pass
    return articles


# ── 文章版本历史 ──────────────────────────────────────────────────────────────

@app.get("/api/articles/{article_id}/versions")
def list_versions(article_id: str):
    ver_dir = BASE_DIR / "data" / "articles" / article_id / "versions"
    if not ver_dir.exists():
        return []
    versions = []
    for f in sorted(ver_dir.glob("*.json"), reverse=True):
        data = json.loads(f.read_text(encoding="utf-8"))
        versions.append({
            "file": f.name,
            "word_count": data.get("word_count", 0),
            "timestamp": data.get("timestamp", ""),
            "preview": data.get("snapshot", ""),
        })
    return versions


@app.get("/api/articles/{article_id}/versions/{filename}")
def get_version(article_id: str, filename: str):
    ver_path = BASE_DIR / "data" / "articles" / article_id / "versions" / filename
    if not ver_path.exists():
        raise HTTPException(404, "版本不存在")
    return json.loads(ver_path.read_text(encoding="utf-8"))


# ── 复制文章 ──────────────────────────────────────────────────────────────────

@app.post("/api/articles/{article_id}/copy")
def copy_article(article_id: str):
    import uuid
    from datetime import datetime, timezone

    src_path = BASE_DIR / "data" / "articles" / f"{article_id}.json"
    if not src_path.exists():
        raise HTTPException(404, "文章不存在")

    src = json.loads(src_path.read_text(encoding="utf-8"))
    new_id = f"art_{uuid.uuid4().hex[:8]}"
    src["id"] = new_id
    src["topic"] = src.get("topic", "") + "（副本）"
    src["created_at"] = datetime.now(timezone.utc).isoformat()

    art_dir = BASE_DIR / "data" / "articles"
    dst_path = art_dir / f"{new_id}.json"
    dst_path.write_text(json.dumps(src, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"ok": True, "id": new_id}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8765)
